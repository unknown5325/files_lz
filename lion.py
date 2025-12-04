import docx  #импортируем библиотеки
import matplotlib.pyplot as plt  #импортируем библиотеки
import pandas as pd  #импортируем библиотеки


text = ''  #создаем пустую строку
tex = ''
wor = []
ks = []
pr = []
words = {}  #создаем словарь

val = []
nuv = []
chlist = [] 

def per_spisok() -> str:
    global text
    doc = docx.Document('lion.docx')  #файл docx
    for paragraph in doc.paragraphs:  #добавляем параграфы в нашу строку
        text += paragraph.text
    return text
def musor():
    global text
    punc = '/?!.,"«»[](){}-–:;—_1234567890xiv'  #убираем лишние символы
    for i in range(0, len(punc)) :
        if punc[i] in text : 
            text = text.replace(punc[i], ' ')
    text = text.lower()
    text = text.split()
    return text
def wordsi(text):
    for item in text:
        if item in words:
            words[item] += 1
        else:
            words[item] = 1
    return words
def proverka(ks, wor,pr):
    for key in words.keys():  #проверка по ключу
        ks.append(words[key])  #столбик с ключами
        wor.append(key)  #столбик со словами
        pr.append(int(words[key])/len(words)*100)  #столбик со значениями
def dat(wor, ks, pr):
    data = {  #создаем таблицу
        'Слово': wor,
        'Частота встречи, раз' : ks,
        'Частота встречи в %' : pr
    }

    df = pd.DataFrame(data)  #создаем именно таблицу
    print(df)  #выводим таблицу
def nov_spisok() -> str:
    global tex
    doc = docx.Document('lion.docx')  #файл docx
    for paragraph in doc.paragraphs:  #добавляем параграфы в нашу строку
        tex += paragraph.text
    return tex
def nov_musor():
    punc = '/?!.,"«»[]()}{-–:;—_1234567890xiv'  #убираем лишиние символы
    for i in range(0, len(punc)) :
        if punc[i] in tex : 
            tex = tex.replace(punc[i], ' ')
    tex = tex.lower()
    return tex
def bykvi():
    ch = {}  #создаем пустой словарь для дальнейшего добавления туда значений и ключа
    for c in tex:   #добавляем в наш пусток ссписок chlist все буквы
        chlist.append(c)
    for item in chlist:
        if item in ch:
            ch[item] += 1
        else:
            ch[item] = 1
    return chlist, ch
def znach_grafik(ch):
    for key in ch.keys():  #подбираем значения для графика
        nuv.append(ch[key])
        val.append(key)
    return nuv, val
def grafik(val, nuv):
    plt.bar(val, nuv)  #создаем график
    plt.xlabel("буквы")  #называем оси
    plt.ylabel("Количество")  #называем оси
    plt.title("Гистограмма количества букв")  #называем график

    plt.show()  #график


chlist = bykvi()
ch = bykvi()
nuv = znach_grafik(ch)
val = znach_grafik(ch)




def main():
    print(ch)
    per_spisok()
    musor()
    wordsi(text)
    proverka(ks, wor, pr)
    dat(wor, ks, pr)
    nov_spisok()
    nov_musor()
    bykvi()
    znach_grafik(ch)
    grafik(val, nuv)


if __name__ == '__main__':
    main()